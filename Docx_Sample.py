#!/usr/bin/env python
# coding: utf-8

# # 0 库引用

# In[1]:


from docx import Document,shared
import random


# # 1 辅助程序

# In[2]:


class _docx:
    def table_reader(table):
        table_list = []
        for i, row in enumerate(table.rows):   # 读每行
            row_content = []
            for cell in row.cells:  # 读一行中的所有单元格
                c = cell.text
                if '\n' in c: c=c.replace('\n','') # 删除换行符
                row_content.append(c)
            table_list.append(row_content)
        return  table_list
    def table_filler(table,table_list):
        for i, row in enumerate(table.rows):    # 读每行
            if i==0: continue                   # 跳过首行
            for j,cell in enumerate(row.cells): # 读一行中的所有单元格
                par = cell.paragraphs[0]
                par.text = str(table_list[i][j]) # 写入内容
                par.style.name = '表内容'        # 添加样式
    def table_transfer(pd_df):
        array = np.array(pd_df.reset_index())
        array_list = array.tolist()
        return array_list
    def table_writer(table,pd_df):
        table_data = _docx.table_transfer(pd_df)
        table_list = _docx.table_reader(table)
        table_list[1:] = table_data
        _docx.table_filler(table,table_list)


# # 2 模板读取与内容查看

# ## 2.1 模板读取

# In[3]:


file = 'Sample_Template.docx'
document = Document(file)


# ## 2.2 模板内容查看

# ### 段落内容

# In[4]:


for i,parag in enumerate(document.paragraphs):
    print('段落{}'.format(i),parag.text)


# ### 表格内容

# In[5]:


table = document.tables[0]


# In[6]:


table_list = _docx.table_reader(table)
table_list


# ### 模板样式查询

# In[7]:


for s in document.styles:
    print(s.type,s.name)


# # 3 写入内容

# In[8]:


i = 8                                                     # 总标题行
document.paragraphs[i].text,i = '程序员种菜指南',i+1       # 写入内容，i=i+1
i += 1                                                    # 再加一行到标题一行

document.paragraphs[i].text,i = '一、程序员为什么种菜',i+1  # 写入内容，i=i+1
document.paragraphs[i].text,i = '1.1、种菜的好处',i+1      # 写入内容，i=i+1
content = '种菜可以吃。'
document.paragraphs[i].text,i = content,i+1               # 写入内容，i=i+1
document.paragraphs[i].text,i = '1.2、程序员种菜的好处',i+1 # 写入内容，i=i+1
content = '程序员可以批量种菜。'
document.paragraphs[i].text,i = content,i+1               # 写入内容，i=i+1

document.paragraphs[i].text,i = '二、种菜流程',i+1         # 写入内容，i=i+1
document.paragraphs[i].text,i = '2.1、种菜流程图',i+1      # 写入内容，i=i+1
content = '程序员种菜主要流程是播种、施肥、浇水和收割。'
document.paragraphs[i].text,i = content,i+1               # 写入内容，i=i+1

picture = '程序员种菜指南.png'
document.paragraphs[i].add_run().add_picture(picture,height=shared.Cm(6),width=shared.Cm(10)) # 在段落内插入图片
i += 1

document.paragraphs[i].text,i = '2.2、种菜成果表',i+1      # 写入内容，i=i+1

for row in range(1,len(table_list)):
    for col in range(1,len(table_list[row])):
        table_list[row][col] = random.randint(0,10)       # 程序员'收集'并填写结果
_docx.table_filler(table,table_list)                      # 将结果写入表格


# # 4 保存文件

# In[9]:


document.save('Sample_Result.docx')


# In[ ]:




